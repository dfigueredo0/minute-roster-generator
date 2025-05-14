import pandas as pd

from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# The following code for handling floating images in a Word document was
# initially reported by user Kill0geR over at the python-docx GitHub page:
# https://github.com/python-openxml/python-docx/issues/159#issuecomment-1955319955
class CT_Anchor(BaseOxmlElement):
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        pic_id = 0
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapSquare wrapText = "bothSides"/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )

def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)

# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page."""
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

def set_cell_background_color(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_document_font(document, font_name):
    styles = document.styles
    for style in styles:
        if style.type == 1:
            style.font.name = font_name

def set_cell_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    for border_name in ['top', 'start', 'bottom', 'end']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)

def set_font(run, font_name='Times New Roman', size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

def set_paragraph_indentation(paragraph, left_indent):
    paragraph.paragraph_format.left_indent = Pt(left_indent)

def insertHR(paragraph, position='bottom'):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '6')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), 'auto')
    pBdr.append(border)

def get_positions(row, delimiter=','):
    if pd.isna(row):
        return []
    return [pos.strip() for pos in row.split(delimiter)]

def add_parliamentary_officers(paragraph, title, role, active_df):
    names = []
    target_roles = [r.strip().lower() for r in role.split(', ')]
    for _, row in active_df.iterrows():
        positions = [p.lower() for p in get_positions(row['Current Office'], '/')]
        if any(r in positions for r in target_roles):
            name = f'{row["First Name"]} {row["Last Name"]}'
            if name not in names:
                names.append(name)
    names_text = ', '.join(names)
    run = paragraph.add_run(f'{title}: {names_text}\n')
    set_font(run)

def add_header(document, header_text, different_header, font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    section = document.sections[0]
    section.different_first_page_header_footer = different_header

    if different_header:
        first_page_header = section.first_page_header
        for paragraph in first_page_header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        if len(first_page_header.paragraphs) == 0:
            first_page_header.add_paragraph()

        first_page_header_paragraph = first_page_header.paragraphs[0]
        first_page_header_run = first_page_header_paragraph.add_run()
        first_page_header_paragraph.alignment = alignment
        set_font(first_page_header_run, font_name, font_size)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(header_text)
    header_paragraph.alignment = alignment
    set_font(header_run, font_name, font_size)

def add_bullet_section(doc, title, bullets):
        p = doc.add_paragraph()
        set_font(p.add_run(title), 'Times New Roman', 11, True)
        for b in bullets:
            bullet = doc.add_paragraph(b, style='List Bullet')
            set_paragraph_indentation(bullet, 36)
    
def set_table_headers(table, headers):
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        if i < len(hdr_cells):
            hdr_cells[i].text = text
    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def apply_table_header_style(cell):
    set_cell_borders(cell)
    set_cell_background_color(cell, '000000')
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor(255, 255, 255)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_table_row(table, texts, r_index, center_cols=None):
    row_cells = table.add_row().cells
    for i, text in enumerate(texts):
        row_cells[i].text = text
    row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'
    for i, cell in enumerate(row_cells):
        set_cell_borders(cell)
        set_cell_background_color(cell, row_color)
        align = WD_ALIGN_PARAGRAPH.CENTER if center_cols and i in center_cols else WD_ALIGN_PARAGRAPH.LEFT
        for paragraph in cell.paragraphs:
            paragraph.alignment = align
    return r_index + 1

def filter_sorted_brothers(active_df, exclude_roles):
    return (active_df[~active_df['Current Office'].str.contains('|'.join(exclude_roles), na=False)]
            .assign(last_lower=active_df['Last Name'].str.lower(),
                    first_lower=active_df['First Name'].str.lower())
            .sort_values(by=['last_lower', 'first_lower'])
            .drop(columns=['last_lower', 'first_lower']))

def get_officers_from_df(df, officer_roles):
      for officer in officer_roles:
        for _, row in df.iterrows():
            positions = [p.strip().lower() for p in get_positions(row['Current Office'], '/')]
            if officer.lower() in positions:
                yield officer, row