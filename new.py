import pandas as pd
import os
import math

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

officers = ['Alpha', 'Beta', 'Pi', 'Iota', 'Sigma', 'Tau', 'Chi', 'Theta One', 'Theta Two', 'Theta Three', 'Upsilon', 'Psi', 'Phi', 'Lambda', 'Asst. Tau', 'Epsilon', 'Gamma']
events = ['Chi', 'Pi', 'Upsilon', 'Psi', 'Phi', 'Gamma', 'Sigma']
advisors = ['Resident Advisor', 'Chapter Advisor', 'Asst. Chapter Advisor']

emDash = u'\u2014'

def set_cell_background_color(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_document_font(document, font_name):
    styles = document.styles
    for style in styles:
        if style.type == 1:
            style.font.name = font_name

def set_cell_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    for border_name in ['top', 'start', 'bottom', 'end']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)

def set_font(run, font_name='Times New Roman', size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

def set_paragraph_indentation(paragraph, left_indent):
    paragraph.paragraph_format.left_indent = Pt(left_indent)

def insertHR(paragraph, position='bottom'):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '6')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), 'auto')
    pBdr.append(border)

def read(excel_file):
    df = pd.read_excel(excel_file, header=1)

    df.columns = df.columns.str.strip()
    df.rename(columns={'Current Office': 'Office'}, inplace=True)
    
    active_df = df[df['Status'] == 'Active'][['Last Name', 'First Name', 'Office']]
    advisor_df = df[df['Office'].isin(advisors)][['Last Name', 'First Name', 'Office']]
    
    active_df['Officers'] = active_df['Office']
    active_df['Roll'] = 'P'
    
    return active_df[['Officers', 'First Name', 'Last Name', 'Roll']], advisor_df

def create_table(doc, columns):
    table = doc.add_table(rows=1, cols=len(columns))
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(columns):
        hdr_cells[i].text = col
    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return table

def fill_table_with_data(table, data, color_even='cccccc', color_odd='FFFFFF'):
    for index, row in data.iterrows():
        row_cells = table.add_row().cells
        if len(row_cells) < len(row):
            raise ValueError("The number of cells in the table row is less than the number of data columns.")
        for i, value in enumerate(row):
            if i < len(row_cells):
                row_cells[i].text = str(value)
        row_color = color_even if index % 2 == 0 else color_odd
        for cell in row_cells:
            set_cell_borders(cell)
            set_cell_background_color(cell, row_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for cell in row_cells[2:]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def create_header(doc, header_text):
    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(header_text)
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

def write_section_title(doc, title, font_name='Times New Roman', size=14, bold=False, hr_position='bottom'):
    paragraph = doc.add_paragraph(title)
    set_font(paragraph.add_run(), font_name, size, bold)
    insertHR(paragraph, hr_position)

def write_paragraph_with_bold_run(doc, text, bold_part, font_name='Times New Roman', size=11):
    paragraph = doc.add_paragraph()
    bold_run = paragraph.add_run(bold_part)
    set_font(bold_run, font_name, size, True)
    paragraph.add_run(text.replace(bold_part, ''))

def create_minutes_doc(docx_output_dir, file_name, header_text, titles, paragraphs, officer_data):
    doc = Document()
    create_header(doc, header_text)
    
    for title in titles:
        write_section_title(doc, title['text'], title.get('font_name', 'Times New Roman'), title.get('size', 14), title.get('bold', False), title.get('hr_position', 'bottom'))

    for paragraph in paragraphs:
        write_paragraph_with_bold_run(doc, paragraph['text'], paragraph['bold_part'], paragraph.get('font_name', 'Times New Roman'), paragraph.get('size', 11))

    officers_table = create_table(doc, ['Officers', 'First Name', 'Last Name', 'Roll'])
    fill_table_with_data(officers_table, officer_data)
    
    doc.save(os.path.join(docx_output_dir, file_name))

def write(active_df, advisor_df):
    docx_output_dir = 'Minutes'
    os.makedirs(docx_output_dir, exist_ok=True)
    xlsx_output_dir = 'Rosters'
    os.makedirs(xlsx_output_dir, exist_ok=True)

    active_df.to_excel(os.path.join(xlsx_output_dir, 'Officer Roster and Minutes Rosters.xlsx'), index=False)

    minutes_config = [
        {
            'file_name': 'Chapter Minutes Outline.docx',
            'header_text': 'Formal Meeting Minutes\n[Date]',
            'titles': [
                {'text': 'Phi Kappa Sigma', 'size': 26, 'bold': True},
                {'text': 'Alpha Epsilon', 'size': 20, 'bold': True},
                {'text': 'Meeting Minutes', 'size': 14, 'bold': False, 'hr_position': 'bottom'},
                {'text': 'Formal Meeting', 'size': 12, 'hr_position': 'top'},
                {'text': 'Date', 'size': 11, 'hr_position': 'bottom'},
                {'text': 'Parliamentary Officers', 'size': 14, 'hr_position': 'top'},
                {'text': 'Chair: Alpha', 'size': 11},
                {'text': 'Secretary: Sigma', 'size': 11},
                {'text': 'Treasurer: Tau', 'size': 11},
                {'text': 'Chaplain: Beta', 'size': 11},
                {'text': 'Sergeants-at-arms: Theta One, Two, Three', 'hr_position': 'bottom'},
                {'text': f'Total active members: {len(active_df)}', 'size': 11},
                {'text': f'Total voting members: {len(active_df)}'},
                {'text': 'Total members in attendance: Attendance', 'size': 11},
                {'text': f'Quorum minimum: {len(active_df) // (3/2)}', 'size': 11},
                {'text': f'Blackball minimum: {math.ceil(len(active_df) * 0.10)}', 'size': 11, 'hr_position': 'bottom'}
            ],
            'paragraphs': [
                {'text': 'Call to Order {0} Time'.format(emDash), 'bold_part': 'Call to Order '},
                {'text': 'Roll -', 'bold_part': 'Roll -'},
                {'text': 'Reading of a section of the Constitution or Chapter By-Laws -', 'bold_part': 'Reading of a section of the Constitution or Chapter By-Laws -'},
                {'text': 'Previous Meeting Minutes -', 'bold_part': 'Previous Meeting Minutes -'},
                {'text': 'Chapter Communications and Letters -', 'bold_part': 'Chapter Communications and Letters -'},
                {'text': 'Proposals and Election of New Members -', 'bold_part': 'Proposals and Election of New Members -'},
                {'text': 'Reports of Officers and Committees -', 'bold_part': 'Reports of Officers and Committees -'},
                {'text': 'Elections -', 'bold_part': 'Elections -'},
                {'text': 'Unfinished Business -', 'bold_part': 'Unfinished Business -'},
                {'text': 'New Business -', 'bold_part': 'New Business -'},
                {'text': 'Comments by the Chapter Advisor, Resident Advisor, and Guests -', 'bold_part': 'Comments by the Chapter Advisor, Resident Advisor, and Guests -'},
                {'text': 'Correction of Minutes -', 'bold_part': 'Correction of Minutes -'},
                {'text': 'Second Roll -', 'bold_part': 'Second Roll -'},
                {'text': 'Closing Comments -', 'bold_part': 'Closing Comments -'},
                {'text': 'Adjournment {0} Time'.format(emDash), 'bold_part': 'Adjournment '},
                {'text': 'Roster {0}'.format(emDash), 'bold_part': 'Roster '}
            ],
            'officer_data': active_df
        },
        {
            'file_name': 'Events Committee Minutes Outline.docx',
            'header_text': 'Events Committee\nXX-XX-XX',
            'titles': [
                {'text': 'Phi Kappa Sigma', 'size': 26, 'bold': True},
                {'text': 'Alpha Epsilon', 'size': 20, 'bold': True},
                {'text': 'Meeting Minutes', 'size': 14, 'bold': False, 'hr_position': 'bottom'},
                {'text': 'Events Committee Meeting', 'size': 12, 'hr_position': 'top'},
                {'text': 'Date', 'size': 11, 'hr_position': 'bottom'},
                {'text': 'Parliamentary Officers', 'size': 14, 'hr_position': 'top'},
                {'text': 'Chair: Chi', 'size': 11},
                {'text': 'Secretary: Sigma', 'size': 11, 'hr_position': 'bottom'}
            ],
            'paragraphs': [
                {'text': 'Call to Order {0} Time'.format(emDash), 'bold_part': 'Call to Order '},
                {'text': 'Unfinished Business', 'bold_part': 'Unfinished Business'},
                {'text': 'Officer Reports', 'bold_part': 'Officer Reports'},
                {'text': 'New Business', 'bold_part': 'New Business'},
                {'text': 'Adjournment {0} Time'.format(emDash), 'bold_part': 'Adjournment '},
                {'text': 'Roster {0}'.format(emDash), 'bold_part': 'Roster '}
            ],
            'officer_data': active_df[active_df['Officers'].str.contains('|'.join(events), na=False)]
        }
    ]

    for config in minutes_config:
        create_minutes_doc(
            docx_output_dir=docx_output_dir,
            file_name=config['file_name'],
            header_text=config['header_text'],
            titles=config['titles'],
            paragraphs=config['paragraphs'],
            officer_data=config['officer_data']
        )

def main():
    excel_file = 'data/AEPKS Roster Fall 2024.xlsx'
    active_df, advisor_df = read(excel_file)
    write(active_df, advisor_df)

if __name__ == "__main__":
    main()
