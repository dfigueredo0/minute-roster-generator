import pandas as pd
import os
import math

from docx import Document
from docx.oxml import parse_xml, register_element_cls, OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.oxml.shape import CT_Picture
from docx.oxml.xmlchemy import BaseOxmlElement, OneAndOnlyOne
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

officers = ['Alpha', 'Beta', 'Pi', 'Iota', 'Sigma', 'Tau', 'Chi', 'Theta One', 'Theta Two', 'Theta Three', 'Upsilon', 'Psi', 'Phi', 'Lambda', 'Asst. Tau', 'Epsilon', 'Gamma']
events = ['Chi', 'Pi', 'Upsilon', 'Psi', 'Phi', 'Gamma', 'Sigma']
exec = ['Alpha', 'Beta', 'Pi', 'Iota', 'Sigma', 'Tau', 'Chi']
advisors = ['Resident Advisor', 'Chapter Advisor', 'Asst. Chapter Advisor']

emDash = u'\u2014'

# The following code for handling floating images in a Word document was
# initially reported by user Kill0geR over at the python-docx GitHub page:
# https://github.com/python-openxml/python-docx/issues/159#issuecomment-1955319955
class CT_Anchor(BaseOxmlElement):
    extent = OneAndOnlyOne('wp:extent')
    docPr = OneAndOnlyOne('wp:docPr')
    graphic = OneAndOnlyOne('a:graphic')

    @classmethod
    def new(cls, cx, cy, shape_id, pic, pos_x, pos_y):
        anchor = parse_xml(cls._anchor_xml(pos_x, pos_y))
        anchor.extent.cx = cx
        anchor.extent.cy = cy
        anchor.docPr.id = shape_id
        anchor.docPr.name = 'Picture %d' % shape_id
        anchor.graphic.graphicData.uri = (
            'http://schemas.openxmlformats.org/drawingml/2006/picture'
        )
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def new_pic_anchor(cls, shape_id, rId, filename, cx, cy, pos_x, pos_y):
        pic_id = 0
        pic = CT_Picture.new(pic_id, filename, rId, cx, cy)
        anchor = cls.new(cx, cy, shape_id, pic, pos_x, pos_y)
        anchor.graphic.graphicData._insert_pic(pic)
        return anchor

    @classmethod
    def _anchor_xml(cls, pos_x, pos_y):
        return (
            '<wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="0" \n'
            '           behindDoc="1" locked="0" layoutInCell="1" allowOverlap="1" \n'
            '           %s>\n'
            '  <wp:simplePos x="0" y="0"/>\n'
            '  <wp:positionH relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionH>\n'
            '  <wp:positionV relativeFrom="page">\n'
            '    <wp:posOffset>%d</wp:posOffset>\n'
            '  </wp:positionV>\n'                    
            '  <wp:extent cx="914400" cy="914400"/>\n'
            '  <wp:wrapSquare wrapText = "bothSides"/>\n'
            '  <wp:docPr id="666" name="unnamed"/>\n'
            '  <wp:cNvGraphicFramePr>\n'
            '    <a:graphicFrameLocks noChangeAspect="1"/>\n'
            '  </wp:cNvGraphicFramePr>\n'
            '  <a:graphic>\n'
            '    <a:graphicData uri="URI not set"/>\n'
            '  </a:graphic>\n'
            '</wp:anchor>' % ( nsdecls('wp', 'a', 'pic', 'r'), int(pos_x), int(pos_y) )
        )

def new_pic_anchor(part, image_descriptor, width, height, pos_x, pos_y):
    rId, image = part.get_or_add_image(image_descriptor)
    cx, cy = image.scaled_dimensions(width, height)
    shape_id, filename = part.next_id, image.filename
    return CT_Anchor.new_pic_anchor(shape_id, rId, filename, cx, cy, pos_x, pos_y)

# refer to docx.text.run.add_picture
def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """Add float picture at fixed position `pos_x` and `pos_y` to the top-left point of page."""
    run = p.add_run()
    anchor = new_pic_anchor(run.part, image_path_or_stream, width, height, pos_x, pos_y)
    run._r.add_drawing(anchor)

def set_cell_background_color(cell, color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color)
    tcPr.append(shd)

def set_document_font(document, font_name):
    styles = document.styles
    for style in styles:
        if style.type == 1:
            style.font.name = font_name

def set_cell_borders(cell):
    tcPr = cell._element.get_or_add_tcPr()
    for border_name in ['top', 'start', 'bottom', 'end']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tcPr.append(border)

def set_font(run, font_name='Times New Roman', size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold

def set_paragraph_indentation(paragraph, left_indent):
    paragraph.paragraph_format.left_indent = Pt(left_indent)

def insertHR(paragraph, position='bottom'):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
        'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
        'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
        'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
        'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
        'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
        'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
        'w:pPrChange'
    )
    border = OxmlElement(f'w:{position}')
    border.set(qn('w:val'), 'single')
    border.set(qn('w:sz'), '6')
    border.set(qn('w:space'), '1')
    border.set(qn('w:color'), 'auto')
    pBdr.append(border)

def get_positions(row, delimiter=','):
    if pd.isna(row):
        return []
    return [pos.strip() for pos in row.split(delimiter)]

def add_parliamentary_officers(paragraph, title, role, active_df):
    names = []
    for r in role.split(', '):
        officer_data = active_df[active_df['Current Office'].apply(lambda x: r in get_positions(x))]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            if name not in names:
                names.append(name)
    names_text = ', '.join(names)
    run = paragraph.add_run(f'{title}: {names_text}\n')
    set_font(run)

def add_header(document, header_text, different_header, font_name='Times New Roman', font_size=11, alignment=WD_ALIGN_PARAGRAPH.RIGHT):
    section = document.sections[0]
    section.different_first_page_header_footer = different_header

    if different_header:
        first_page_header = section.first_page_header
        for paragraph in first_page_header.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        if len(first_page_header.paragraphs) == 0:
            first_page_header.add_paragraph()

        first_page_header_paragraph = first_page_header.paragraphs[0]
        first_page_header_run = first_page_header_paragraph.add_run()
        first_page_header_paragraph.alignment = alignment
        set_font(first_page_header_run, font_name, font_size)

    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run(header_text)
    header_paragraph.alignment = alignment
    set_font(header_run, font_name, font_size)

def create_chapter_minutes(docx_output_dir, active_df, advisor_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    add_header(doc, 'Formal Meeting Minutes\nDate', True)

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_CREST.png', width=Inches(7), height=Inches(9), pos_x=Pt(-225), pos_y=Pt(80))

    title_paragraph = doc.add_paragraph()
    set_font(title_paragraph.add_run('Phi Kappa Sigma\n'), size=26, bold=True)
    set_font(title_paragraph.add_run('Alpha Epsilon\n'), size=20, bold=True)
    set_font(title_paragraph.add_run('Meeting Minutes'), size=14)
    insertHR(title_paragraph)

    meeting = doc.add_paragraph()
    set_font(meeting.add_run('Formal Meeting\n'), size=12)
    insertHR(meeting, position='top')
    set_font(meeting.add_run('Date'))

    parliamentary_officers = doc.add_paragraph()
    set_font(parliamentary_officers.add_run('Parliamentary Officers\n'), size=14)
    insertHR(parliamentary_officers, position='top')
    roles = [('Chair', 'Alpha'), ('Secretary', 'Sigma'), ('Treasurer', 'Tau'), ('Chaplain', 'Beta'), ('Sergeants-at-arms', 'Theta One, Theta Two, Theta Three')]

    for title, role in roles:
        add_parliamentary_officers(parliamentary_officers, title, role, active_df)

    insertHR(parliamentary_officers)

    active_members = doc.add_paragraph()
    set_font(active_members.add_run(f'Total active members: {len(active_df)}\n'))
    set_font(active_members.add_run(f'Total voting members: {len(active_df)}\n'))
    set_font(active_members.add_run('Total members in attendance: Attendance\n'))
    set_font(active_members.add_run(f'Quorum minimum {int(len(active_df) // (3/2))}\n'))
    set_font(active_members.add_run(f'Blackball minimum: {math.ceil(len(active_df) * 0.10)} \t(10%)\n'))
    insertHR(active_members)

    doc.add_page_break()

    call = doc.add_paragraph()
    set_font(call.add_run('Call to Order - Time'), font_name='Helvetica Neue', bold=True)

    roll = doc.add_paragraph()
    set_font(roll.add_run('Roll -'), font_name='Helvetica Neue', bold=True)

    reading = doc.add_paragraph()
    set_font(reading.add_run('Reading of a section of the Constitution or Chapter By-Laws -'), font_name='Helvetica Neue', bold=True)

    previous = doc.add_paragraph()
    set_font(previous.add_run('Previous Meeting Minutes -'), font_name='Helvetica Neue', bold=True)

    chapter_comms = doc.add_paragraph()
    set_font(chapter_comms.add_run('Chapter Communications and Letters -'), font_name='Helvetica Neue', bold=True)

    proposals = doc.add_paragraph()
    set_font(proposals.add_run('Proposals and Election of New Members -'), font_name='Helvetica Neue', bold=True)
    nms = {
        'NM1': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:'],
        'NM2': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:'],
        'NM3': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:']
    }

    for nm, sub_items in nms.items():
        bullet_point = doc.add_paragraph(style='List Bullet')
        set_font(bullet_point.add_run(nm), font_name='Helvetica Neue')
        set_paragraph_indentation(bullet_point, 36)
        
        for sub_item in sub_items:
            sub_bullet_point = doc.add_paragraph(style='List Bullet 2')
            set_font(sub_bullet_point.add_run(sub_item), font_name='Helvetica Neue')
            set_paragraph_indentation(sub_bullet_point, 72)

    reports = doc.add_paragraph()
    set_font(reports.add_run('Reports of Officers and Committees -'), font_name='Helvetica Neue', bold=True)
    officer_committees = [
        'Executive Committee', 'Finance Committee', 'Recuirment Committee', 
        'Events Committee', 'Internal Operations Committee', 'Bylaws Committee', 
        'Greek Council Liasion (Gamma)', 'Epsilon', 'Lambda', 'Asst. Tau', 'Phi', 
        'Psi', 'Upsilon', 'Theta Three', 'Theta Two', 'Theta One', 'Chi', 'Tau',
        'Sigma', 'Iota', 'Pi', 'Beta', 'Alpha'
    ]

    for oc in officer_committees:
        bullet_point = doc.add_paragraph(style='List Bullet')
        set_font(bullet_point.add_run(oc), font_name='Helvetica Neue')
        set_paragraph_indentation(bullet_point, 36)

    elections = doc.add_paragraph()
    set_font(elections.add_run('Elections -'), font_name='Helvetica Neue', bold=True)

    unfinished = doc.add_paragraph()
    set_font(unfinished.add_run('Unfinished Business -'), font_name='Helvetica Neue', bold=True)

    new_business = doc.add_paragraph()
    set_font(new_business.add_run('New Business -'), font_name='Helvetica Neue', bold=True)

    comments = doc.add_paragraph()
    set_font(comments.add_run('Comments by the Chapter Advisor, Resident Advisor, and Guests -'), font_name='Helvetica Neue', bold=True)

    corrections = doc.add_paragraph()
    set_font(corrections.add_run('Correction of Minutes -'), font_name='Helvetica Neue', bold=True)

    roll2 = doc.add_paragraph()
    set_font(roll2.add_run('Second Roll -'), font_name='Helvetica Neue', bold=True)

    closing = doc.add_paragraph()
    set_font(closing.add_run('Closing Comments -'), font_name='Helvetica Neue', bold=True)
    bullet_point = doc.add_paragraph(style='List Bullet')
    set_font(bullet_point.add_run('Announcements'), font_name='Helvetica Neue')
    set_paragraph_indentation(bullet_point, 36)
    sub_bullet_point = doc.add_paragraph(style='List Bullet 2')
    set_font(sub_bullet_point.add_run(''), font_name='Helvetica Neue')
    set_paragraph_indentation(sub_bullet_point, 72)

    bullet_point1 = doc.add_paragraph(style='List Bullet')
    set_font(bullet_point1.add_run('Betterment'), font_name='Helvetica Neue')
    set_paragraph_indentation(bullet_point1, 36)
    sub_bullet_point1 = doc.add_paragraph(style='List Bullet 2')
    set_font(sub_bullet_point1.add_run(''), font_name='Helvetica Neue')
    set_paragraph_indentation(sub_bullet_point1, 72)

    adjourn = doc.add_paragraph()
    set_font(adjourn.add_run('Adjournment - Time'), font_name='Helvetica Neue', bold=True)

    doc.add_page_break()

    roster = doc.add_paragraph()
    set_font(roster.add_run('Roster -'), font_name='Helvetica Neue', bold=True)

    set_document_font(document=doc, font_name='Calibri')
    officers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'

    hdr_cells[0].merge(hdr_cells[1])
  
    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    for officer in officers:
        officer_data = active_df[active_df['Current Office'].apply(lambda x: officer in get_positions(x))]
        for index, row in officer_data.iterrows():
            row_cells = officers_table.add_row().cells
            row_cells[0].text = officer
            row_cells[1].text = f'{row["First Name"]} {row["Last Name"]}'
            row_cells[2].text = 'P'
            row_cells[3].text = 'P'

            row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

            for cell in row_cells:
                set_cell_borders(cell)
                set_cell_background_color(cell, row_color)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for cell in [row_cells[2], row_cells[3]]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            r_index += 1

    brothers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Brothers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'

    hdr_cells[0].merge(hdr_cells[1])

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    brothers_data = active_df[~active_df['Current Office'].str.contains('|'.join(officers), na=False)]
    for index, row in brothers_data.iterrows():
        row_cells = brothers_table.add_row().cells
        row_cells[0].text = f'{row["Last Name"]}'
        row_cells[1].text = f'{row["First Name"]}'
        row_cells[2].text = 'P'
        row_cells[3].text = 'P'

        row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

        for cell in row_cells:
            set_cell_borders(cell)
            set_cell_background_color(cell, row_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for cell in [row_cells[2], row_cells[3]]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_index += 1

    advisor_table = doc.add_table(rows=1, cols=4)
    hdr_cells = advisor_table.rows[0].cells
    hdr_cells[0].text = 'Role'
    hdr_cells[1].text = 'Chapter Staff'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    for advisor in advisors:
        advisor_data = advisor_df[advisor_df['Current Office'] == advisor]
        for index, row in advisor_data.iterrows():
            row_cells = advisor_table.add_row().cells
            row_cells[0].text = advisor
            row_cells[1].text = f'{row["First Name"]} {row["Last Name"]}'
            row_cells[2].text = 'E' if advisor in ['Chapter Advisor', 'Asst. Chapter Advisor'] else 'P'
            row_cells[3].text = 'E' if advisor in ['Chapter Advisor', 'Asst. Chapter Advisor'] else 'P'

            row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

            for cell in row_cells:
                set_cell_borders(cell)
                set_cell_background_color(cell, row_color)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

            for cell in [row_cells[2], row_cells[3]]:
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            r_index += 1

    doc.save(os.path.join(docx_output_dir, 'Chapter Minutes Outline.docx'))

def create_bylaws_minutes(docx_output_dir, active_df):
    doc = Document()

    add_header(doc,'Bylaws Committee Meeting\nXX-XX-XX', False)

    title = doc.add_paragraph()
    set_font(title.add_run('Phi Kappa Sigma\n'), 'Times New Roman', 26, True)
    set_font(title.add_run('Alpha Epsilon\n'), 'Times New Roman', 20, True)
    set_font(title.add_run('Meeting Minutes'), 'Times New Roman', 14)
    insertHR(title)

    bylaws = doc.add_paragraph()
    set_font(bylaws.add_run('Bylaws Committee Meeting\n'), 'Times New Roman', 12)
    insertHR(bylaws, 'top')
    set_font(bylaws.add_run('Date'), 'Times New Roman', 11)
    insertHR(bylaws)

    parliamentary_officers = doc.add_paragraph()
    set_font(parliamentary_officers.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)

    roles = [('Chair', 'Sigma'), ('Secretary', 'Sigma')]

    for title, role in roles:
        add_parliamentary_officers(parliamentary_officers, title, role, active_df)

    insertHR(parliamentary_officers)

    call_to_order_paragraph = doc.add_paragraph()
    call_to_order_run = call_to_order_paragraph.add_run('Call to Order ')
    set_font(call_to_order_run, 'Times New Roman', 11, bold=True)
    call_to_order_paragraph.add_run('{0} Time'.format(emDash))

    text = doc.add_paragraph('Unfinished Business')
    set_font(text.add_run(), 'Times New Roman', 11, True)
    
    bullet_point = doc.add_paragraph('None', style='List Bullet')
    set_paragraph_indentation(bullet_point, 36)

    text = doc.add_paragraph('New Business')
    set_font(text.add_run(), 'Times New Roman', 11, True)

    adjournment_paragraph = doc.add_paragraph()
    adjournment_run = adjournment_paragraph.add_run('Adjournment ')
    set_font(adjournment_run, 'Times New Roman', 11, True)
    adjournment_paragraph.add_run('{0} Time'.format(emDash))
    
    roster_paragraph = doc.add_paragraph()
    roster_run = roster_paragraph.add_run('Roster ')
    set_font(roster_run, 'Times New Roman', 11, True)
    roster_paragraph.add_run('{0}'.format(emDash))

    set_document_font(document=doc, font_name='Calibri')

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    positions = ['Chair', 'Secretary']
    for position in positions:
        officer_data = active_df[active_df['Current Office'].str.contains('Sigma', na=False)]
        if not officer_data.empty:
            for index, row in officer_data.iterrows():
                row_cells = officers_table.add_row().cells
                row_cells[0].text = position
                row_cells[1].text = f'{row["First Name"]} {row["Last Name"]}'
                row_cells[2].text = 'P'

                row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

                for cell in row_cells:
                    set_cell_borders(cell)
                    set_cell_background_color(cell, row_color)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for cell in [row_cells[2]]:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                r_index += 1

    others_table = doc.add_table(rows=1, cols=3)
    hdr_cells = others_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(officers), na=False)]
    for index, row in others_data.iterrows():
        row_cells = others_table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = ''
        row_cells[2].text = 'P'

        row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

        for cell in row_cells:
            set_cell_borders(cell)
            set_cell_background_color(cell, row_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for cell in [row_cells[2]]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_index += 1

    doc.save(os.path.join(docx_output_dir, 'Bylaws Committe Minutes Outline.docx'))

def create_events_minutes(docx_output_dir, active_df):
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('Events Committee\nXX-XX-XX')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    title = doc.add_paragraph('Phi Kappa Sigma')
    set_font(title.add_run(), 'Times New Roman', 26, True)

    subtitle = doc.add_paragraph('Alpha Epsilon')
    set_font(subtitle.add_run(), 'Times New Roman', 20, True)

    meeting_title = doc.add_paragraph('Meeting Minutes')
    set_font(meeting_title.add_run(), 'Times New Roman', 14)
    insertHR(meeting_title)

    bylaws = doc.add_paragraph('Bylaws Committee Meeting')
    set_font(bylaws.add_run(), 'Times New Roman', 12)
    insertHR(bylaws, 'top')

    date = doc.add_paragraph('Date')
    set_font(date.add_run(), 'Times New Roman', 11)
    insertHR(date)

    parliamentary_officers = doc.add_paragraph('Parliamentary Officers')
    set_font(parliamentary_officers.add_run(), 'Times New Roman', 14)
    insertHR(parliamentary_officers, 'top')

    chair = doc.add_paragraph('Chair: Chi')
    set_font(chair.add_run(), 'Times New Roman', 11)

    secretary = doc.add_paragraph('Secretary: Sigma')
    set_font(secretary.add_run(), 'Times New Roman', 11)
    insertHR(secretary)

    call_to_order_paragraph = doc.add_paragraph()
    call_to_order_run = call_to_order_paragraph.add_run('Call to Order ')
    set_font(call_to_order_run, 'Times New Roman', 11, bold=True)
    call_to_order_paragraph.add_run('{0} Time'.format(emDash))

    text = doc.add_paragraph('Unfinished Business')
    set_font(text.add_run(), 'Times New Roman', 11, True)
    
    text = doc.add_paragraph('Officer Reports')
    set_font(text.add_run(), 'Times New Roman', 11, True)

    chi_bullet_point = doc.add_paragraph('Chi', style='List Bullet')
    set_paragraph_indentation(chi_bullet_point, 36)

    pi_bullet_point = doc.add_paragraph('Pi', style='List Bullet')
    set_paragraph_indentation(pi_bullet_point, 36)

    upsilon_bullet_point = doc.add_paragraph('Upsilon', style='List Bullet')
    set_paragraph_indentation(upsilon_bullet_point, 36)

    psi_bullet_point = doc.add_paragraph('Psi', style='List Bullet')
    set_paragraph_indentation(psi_bullet_point, 36)

    phi_bullet_point = doc.add_paragraph('Phi', style='List Bullet')
    set_paragraph_indentation(phi_bullet_point, 36)

    gamma_bullet_point = doc.add_paragraph('Gamma', style='List Bullet')
    set_paragraph_indentation(gamma_bullet_point, 36)

    text = doc.add_paragraph('New Business')
    set_font(text.add_run(), 'Times New Roman', 11, True)
    bullet_point = doc.add_paragraph('', style='List Bullet')
    set_paragraph_indentation(bullet_point, 36)

    adjournment_paragraph = doc.add_paragraph()
    adjournment_run = adjournment_paragraph.add_run('Adjournment ')
    set_font(adjournment_run, 'Times New Roman', 11, True)
    adjournment_paragraph.add_run('{0} Time'.format(emDash))
    
    roster_paragraph = doc.add_paragraph()
    roster_run = roster_paragraph.add_run('Roster ')
    set_font(roster_run, 'Times New Roman', 11, True)
    roster_paragraph.add_run('{0}'.format(emDash))

    set_document_font(document=doc, font_name='Calibri')

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    for event in events:
        officer_data = active_df[active_df['Current Office'].str.contains(event, na=False)]
        if not officer_data.empty:
            for index, row in officer_data.iterrows():
                row_cells = officers_table.add_row().cells
                row_cells[0].text = event
                row_cells[1].text = f'{row["First Name"]} {row["Last Name"]}'
                row_cells[2].text = 'P'

                row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

                for cell in row_cells:
                    set_cell_borders(cell)
                    set_cell_background_color(cell, row_color)
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT

                for cell in [row_cells[2]]:
                    for paragraph in cell.paragraphs:
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                r_index += 1

    others_table = doc.add_table(rows=1, cols=3)
    hdr_cells = others_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])

    for cell in hdr_cells:
        set_cell_borders(cell)
        set_cell_background_color(cell, '000000')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(officers), na=False)]
    for index, row in others_data.iterrows():
        row_cells = others_table.add_row().cells
        row_cells[0].text = ''
        row_cells[1].text = ''
        row_cells[2].text = 'P'

        row_color = 'cccccc' if r_index % 2 == 0 else 'FFFFFF'

        for cell in row_cells:
            set_cell_borders(cell)
            set_cell_background_color(cell, row_color)
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        for cell in [row_cells[2]]:
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        r_index += 1

    doc.save(os.path.join(docx_output_dir, 'Events Committe Minutes Outline.docx'))

def create_exec_minutes(docx_output_dir, active_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('Executive Council Meeting Minutes\nDate')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_INSIGNIA.png', width=Inches(3.65), height=Inches(3.95), pos_x=0, pos_y=0)

    title = doc.add_paragraph('Phi Kappa Sigma')
    set_font(title.add_run(), 'Times New Roman', 26, True)

    subtitle = doc.add_paragraph('Alpha Epsilon')
    set_font(subtitle.add_run(), 'Times New Roman', 20, True)

    meeting_title = doc.add_paragraph('Meeting Minutes')
    set_font(meeting_title.add_run(), 'Times New Roman', 14)
    insertHR(meeting_title)

    bylaws = doc.add_paragraph('Bylaws Committee Meeting')
    set_font(bylaws.add_run(), 'Times New Roman', 12)
    insertHR(bylaws, 'top')

    date = doc.add_paragraph('Date')
    set_font(date.add_run(), 'Times New Roman', 11)
    insertHR(date)

    parliamentary_officers = doc.add_paragraph('Parliamentary Officers')
    set_font(parliamentary_officers.add_run(), 'Times New Roman', 14)
    insertHR(parliamentary_officers, 'top')

    chair = doc.add_paragraph('Chair: Alpha')
    set_font(chair.add_run(), 'Times New Roman', 11)

    secretary = doc.add_paragraph('Secretary: Sigma')
    set_font(secretary.add_run(), 'Times New Roman', 11)

    doc.save(os.path.join(docx_output_dir, 'Exec Minutes Outline.docx'))

def create_finance_minutes(docx_output_dir, active_df):
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('Finance Committee Meeting\nXX-XX-XX')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    doc.save(os.path.join(docx_output_dir, 'Finance Committee Outline.docx'))

def create_house_minutes(docx_output_dir, active_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('House Meeting Minutes\nDate')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_INSIGNIA.png', width=Inches(3.65), height=Inches(3.95), pos_x=Pt(5), pos_y=Pt(92))

    title = doc.add_paragraph('Phi Kappa Sigma')
    set_font(title.add_run(), 'Times New Roman', 26, True)

    subtitle = doc.add_paragraph('Alpha Epsilon')
    set_font(subtitle.add_run(), 'Times New Roman', 20, True)

    meeting_title = doc.add_paragraph('Meeting Minutes')
    set_font(meeting_title.add_run(), 'Times New Roman', 14)
    insertHR(meeting_title)

    bylaws = doc.add_paragraph('Bylaws Committee Meeting')
    set_font(bylaws.add_run(), 'Times New Roman', 12)
    insertHR(bylaws, 'top')

    date = doc.add_paragraph('Date')
    set_font(date.add_run(), 'Times New Roman', 11)
    insertHR(date)

    parliamentary_officers = doc.add_paragraph('Parliamentary Officers')
    set_font(parliamentary_officers.add_run(), 'Times New Roman', 14)
    insertHR(parliamentary_officers, 'top')

    chair = doc.add_paragraph('Chair: Alpha')
    set_font(chair.add_run(), 'Times New Roman', 11)

    secretary = doc.add_paragraph('Secretary: Sigma')
    set_font(secretary.add_run(), 'Times New Roman', 11)

    doc.save(os.path.join(docx_output_dir, 'House Minutes Outline.docx'))

def create_IOC_minutes(docx_output_dir, active_df):
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('Internal Operations Committee\nXX-XX-XX')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    doc.save(os.path.join(docx_output_dir, 'IOC Minutes Outline.docx'))

def create_jboard_minutes(docx_output_dir, active_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    section = doc.sections[0]
    header = section.header
    header_paragraph = header.paragraphs[0]
    header_run = header_paragraph.add_run('Judicial Board Meeting Minutes\nXX-XX-XX')
    header_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_font(header_run, 'Times New Roman', 11)

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_FAST_F.png', width=Inches(3.25), height=Inches(5.5), pos_x=0, pos_y=0)

    doc.save(os.path.join(docx_output_dir, 'Judical Board Meeting Minutes Outline.docx'))

def create_roster(xlsx_output_dir, active_df, advisor_df):
    active_df.to_excel(os.path.join(xlsx_output_dir, 'Officer Roster and Minutes Rosters.xlsx'), index=False)

def read(excel_file):
    df = pd.read_excel(excel_file, header=1)
    active_df = df[df['Status'] == 'Active'][['Last Name', 'First Name', 'Current Office']]
    advisor_df = df[df['Current Office'].isin(advisors)][['Last Name', 'First Name', 'Current Office']]

    return active_df, advisor_df

def write(active_df, advisor_df):
    docx_output_dir = 'Minutes'
    os.makedirs(docx_output_dir, exist_ok=True)
    xlsx_output_dir = 'Rosters'
    os.makedirs(xlsx_output_dir, exist_ok=True)

    create_roster(xlsx_output_dir, active_df, advisor_df)

    create_chapter_minutes(docx_output_dir, active_df, advisor_df)

    create_bylaws_minutes(docx_output_dir, active_df)

    create_events_minutes(docx_output_dir, active_df)

    create_exec_minutes(docx_output_dir, active_df)

    create_finance_minutes(docx_output_dir, active_df)

    create_house_minutes(docx_output_dir, active_df)

    create_IOC_minutes(docx_output_dir, active_df)

    create_jboard_minutes(docx_output_dir, active_df)

def main():
    excel_file = 'data/AEPKS Roster Fall 2024.xlsx'
    active_df, advisor_df = read(excel_file)
    write(active_df, advisor_df)

if __name__ == "__main__":
    main()
