import os
import math

from docx import Document
from docx.oxml import register_element_cls, OxmlElement
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from utils import *
from constants import *

def create_bylaws_minutes(docx_output_dir, active_df):
    doc = Document()
    add_header(doc, 'Bylaws Committee Meeting\nXX-XX-XX', False)

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), 'Times New Roman', size, bold)
    insertHR(title)

    bylaws = doc.add_paragraph()
    set_font(bylaws.add_run('Bylaws Committee Meeting\n'), 'Times New Roman', 12)
    insertHR(bylaws, 'top')
    set_font(bylaws.add_run('Date'), 'Times New Roman', 11)
    insertHR(bylaws)

    parliamentary_officers = doc.add_paragraph()
    set_font(parliamentary_officers.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)
    roles = [('Chair', 'Sigma'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officers, title, role, active_df)
    insertHR(parliamentary_officers)

    add_bullet_section(doc, 'Unfinished Business', ['None'])
    add_bullet_section(doc, 'New Business', [''])

    adj = doc.add_paragraph()
    set_font(adj.add_run(f'Adjournment {emDash} Time'), 'Times New Roman', 11, True)

    doc.add_page_break()

    roster = doc.add_paragraph()
    set_font(roster.add_run(f'Roster {emDash}'), 'Times New Roman', 11, True)

    set_document_font(doc, font_name='Calibri')

    table_title = doc.add_table(rows=1, cols=1)
    cell = table_title.cell(0, 0)
    cell.text = 'BYLAWS COMMITTEE'
    cell.paragraphs[0].style = 'No Spacing'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Roll'

    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    for position in ['Chair', 'Secretary']:
        officer_data = active_df[active_df['Current Office'].str.contains('Sigma', na=False)]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            r_index = add_table_row(officers_table, [position, name, 'P'], r_index, center_cols=[2])

    brothers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'

    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(officers), na=False)]
    for _, row in others_data.iterrows():
        if r_index >= 5:
            break
        r_index = add_table_row(brothers_table, ['', '', 'P'], r_index, center_cols=[2])

    doc.save(os.path.join(docx_output_dir, 'Bylaws Committe Minutes Outline.docx'))

def create_chapter_minutes(docx_output_dir, active_df, advisor_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()
    add_header(doc, 'Formal Meeting Minutes\nDate', True)

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_CREST.png', width=Inches(7), height=Inches(9), pos_x=Pt(-225), pos_y=Pt(80))

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), size=size, bold=bold)
    insertHR(title)

    meeting = doc.add_paragraph()
    set_font(meeting.add_run('Formal Meeting\n'), size=12)
    insertHR(meeting, position='top')
    set_font(meeting.add_run('Date'))

    roles = [('Chair', 'Alpha'), ('Secretary', 'Sigma'), ('Treasurer', 'Tau'), 
             ('Chaplain', 'Beta'), ('Sergeants-at-Arms', 'Theta One, Theta Two, Theta Three')]
    parliamentary = doc.add_paragraph()
    set_font(parliamentary.add_run('Parliamentary Officers\n'), size=14)
    insertHR(parliamentary, position='top')
    for title, role in roles:
        add_parliamentary_officers(parliamentary, title, role, active_df)
    insertHR(parliamentary)

    stats = doc.add_paragraph()
    num_members = len(active_df)
    quorum = int(num_members // (3/2))
    blackball = math.ceil(num_members * 0.10)
    for line in [
        f'Total active members: {num_members}\n',
        f'Total voting members: {num_members}\n',
        'Total members in attendance: Attendance\n',
        f'Quorum minimum {quorum}\n',
        f'Blackball minimum: {blackball} \t(10%)\n'
    ]:
        set_font(stats.add_run(line))
    insertHR(stats)

    doc.add_page_break()

    for header in [
        'Call to Order - Time', 'Roll -', 
        'Reading of a section of the Constitution or Chapter By-Laws -',
        'Previous Meeting Minutes -', 'Chapter Communications and Letters -',
        'Proposals and Election of New Members -'
    ]:
        p = doc.add_paragraph()
        set_font(p.add_run(header), font_name='Helvetica Neue', bold=True)

    nms = {
        'NM1': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:'],
        'NM2': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:'],
        'NM3': ['Pi:', 'Iota:', 'PF:', '+:', '-:', 'C:', 'Pass:']
    }
    for nm, items in nms.items():
        bullet = doc.add_paragraph(style='List Bullet')
        set_font(bullet.add_run(nm), font_name='Helvetica Neue')
        set_paragraph_indentation(bullet, 36)
        for sub in items:
            sub_bullet = doc.add_paragraph(style='List Bullet 2')
            set_font(sub_bullet.add_run(sub), font_name='Helvetica Neue')
            set_paragraph_indentation(sub_bullet, 72)

    report_titles = ['Reports of Officers and Committees -']
    committees = [
        'Executive Committee', 'Finance Committee', 'Recuirment Committee', 
        'Events Committee', 'Internal Operations Committee', 'Bylaws Committee', 
        'Greek Council Liasion (Gamma)', 'Epsilon', 'Lambda', 'Asst. Tau', 'Phi', 
        'Psi', 'Upsilon', 'Theta Three', 'Theta Two', 'Theta One', 'Chi', 'Tau',
        'Sigma', 'Iota', 'Pi', 'Beta', 'Alpha'
    ]
    for title in report_titles:
        p = doc.add_paragraph()
        set_font(p.add_run(title), font_name='Helvetica Neue', bold=True)
    for c in committees:
        bullet = doc.add_paragraph(style='List Bullet')
        set_font(bullet.add_run(c), font_name='Helvetica Neue')
        set_paragraph_indentation(bullet, 36)

    for section in [
        'Elections -', 'Unfinished Business -', 'New Business -',
        'Comments by the Chapter Advisor, Resident Advisor, and Guests -',
        'Correction of Minutes -', 'Second Roll -', 'Closing Comments -'
    ]:
        p = doc.add_paragraph()
        set_font(p.add_run(section), font_name='Helvetica Neue', bold=True)

    for header in ['Announcements', 'Betterment']:
        bullet = doc.add_paragraph(style='List Bullet')
        set_font(bullet.add_run(header), font_name='Helvetica Neue')
        set_paragraph_indentation(bullet, 36)
        sub = doc.add_paragraph(style='List Bullet 2')
        set_font(sub.add_run(''), font_name='Helvetica Neue')
        set_paragraph_indentation(sub, 72)

    adj = doc.add_paragraph()
    set_font(adj.add_run('Adjournment - Time'), font_name='Helvetica Neue', bold=True)

    doc.add_page_break()

    roster = doc.add_paragraph()
    set_font(roster.add_run('Roster -'), font_name='Helvetica Neue', bold=True)
    set_document_font(doc, font_name='Calibri')

    officers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'

    hdr_cells[0].merge(hdr_cells[1])
    
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    for officer, row in get_officers_from_df(active_df, officers):
        name = f'{row["First Name"]} {row["Last Name"]}'
        r_index = add_table_row(officers_table, [officer, name, 'P', 'P'], r_index, center_cols=[2, 3])

    brothers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Brothers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'

    hdr_cells[0].merge(hdr_cells[1])
    
    for cell in hdr_cells:
        apply_table_header_style(cell)

    brothers_data = filter_sorted_brothers(active_df, officers)
    r_index = 0
    for _, row in brothers_data.iterrows():
        r_index = add_table_row(brothers_table, [row["Last Name"], row["First Name"], 'P', 'P'], r_index, center_cols=[2, 3])

    advisor_table = doc.add_table(rows=1, cols=4)
    set_table_headers(advisor_table, ['Role', 'Chapter Staff', 'Opening Roll', 'Closing Roll'])

    r_index = 0
    for advisor in advisors:
        advisor_data = advisor_df[advisor_df['Current Office'] == advisor]
        for _, row in advisor_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            symbol = 'E' if advisor in ['Chapter Advisor', 'Asst. Chapter Advisor'] else 'P'
            r_index = add_table_row(advisor_table, [advisor, name, symbol, symbol], r_index, center_cols=[2, 3])

    doc.save(os.path.join(docx_output_dir, 'Chapter Minutes Outline.docx'))

def create_events_minutes(docx_output_dir, active_df):
    doc = Document()
    add_header(doc, 'Events Committee\nXX-XX-XX', False)

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), 'Times New Roman', size, bold)
    insertHR(title)

    events_p = doc.add_paragraph()
    set_font(events_p.add_run('Events Committee\n'), 'Times New Roman', 12)
    insertHR(events_p, 'top')
    set_font(events_p.add_run('Date'), 'Times New Roman', 11)
    insertHR(events_p)

    parliamentary_officer = doc.add_paragraph()
    set_font(parliamentary_officer.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)
    roles = [('Chair', 'Chi'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officer, title, role, active_df)
    insertHR(parliamentary_officer)

    call = doc.add_paragraph()
    set_font(call.add_run('Call to Order '), 'Times New Roman', 11, bold=True)
    call.add_run(f'{emDash} Time')

    add_bullet_section(doc, 'Unfinished Business', ['None'])

    add_bullet_section(doc, 'Officer Reports', ['Chi', 'Pi', 'Upsilon', 'Psi', 'Phi', 'Gamma'])

    add_bullet_section(doc, 'New Business', [''])

    adj = doc.add_paragraph()
    set_font(adj.add_run('Adjournment '), 'Times New Roman', 11, True)
    adj.add_run(f'{emDash} Time')

    doc.add_page_break()

    roster = doc.add_paragraph()
    set_font(roster.add_run('Roster '), 'Times New Roman', 11, True)
    roster.add_run(f'{emDash}')

    set_document_font(document=doc, font_name='Calibri')

    table_title = doc.add_table(rows=1, cols=1)
    cell = table_title.cell(0, 0)
    cell.text = 'EVENTS COMMITTEE'
    cell.paragraphs[0].style = 'No Spacing'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Brothers'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    for role in events:
        officer_data = active_df[active_df['Current Office'].str.contains(role, na=False)]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            r_index = add_table_row(officers_table, [role, name, 'P'], r_index, center_cols=[2])

    brothers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)
        
    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(officers), na=False)]
    for _, row in others_data.iterrows():
        if r_index >= 5:
            break
        r_index = add_table_row(brothers_table, ['', '', 'P'], r_index, center_cols=[2])

    doc.save(os.path.join(docx_output_dir, 'Events Committe Minutes Outline.docx'))

def create_exec_minutes(docx_output_dir, active_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_FAST_F.png', width=Inches(1.6), height=Inches(2.75), pos_x=Pt(90), pos_y=Pt(70))

    title_paragraph = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title_paragraph.add_run(text), size=size, bold=bold)
    insertHR(title_paragraph)

    meeting = doc.add_paragraph()
    set_font(meeting.add_run('Executive Council Minutes\n'), 'Times New Roman', 12)
    insertHR(meeting, 'top')
    set_font(meeting.add_run('Date'), 'Times New Roman', 11)
    insertHR(meeting)

    parliamentary_officer = doc.add_paragraph('Parliamentary Officers\n')
    set_font(parliamentary_officer.add_run(), 'Times New Roman', 14)
    roles = [('Chair', 'Alpha'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officer, title, role, active_df)
    insertHR(parliamentary_officer)

    set_font(doc.add_paragraph().add_run('Call to Order - Time'), 'Times New Roman', 11, True)
    add_bullet_section(doc, f'Officer Reports {emDash}', ['Alpha', 'Beta', 'Pi', 'Sigma', 'Tau', 'Chi', 'Iota'])
    set_font(doc.add_paragraph().add_run(f'Business {emDash}'), 'Times New Roman', 11, True)
    set_font(doc.add_paragraph().add_run(f'Adjournment {emDash}'), 'Times New Roman', 11, True)

    doc.add_page_break()

    set_font(doc.add_paragraph().add_run(f'Roster {emDash}'), 'Times New Roman', 11, True)

    title_table = doc.add_table(rows=1, cols=1)
    title_cell = title_table.cell(0, 0)
    title_cell.text = 'EXECUTIVE COUNCIL MEETING'
    title_cell.paragraphs[0].style = 'No Spacing'
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    officers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    for officer in exec:
        officer_data = active_df[active_df['Current Office'].str.match(officer, na=False)]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            r_index = add_table_row(officers_table, [officer, name, 'P', 'P'], r_index, center_cols=[2, 3])

    doc.save(os.path.join(docx_output_dir, 'Exec Minutes Outline.docx'))

def create_finance_minutes(docx_output_dir, active_df):
    doc = Document()
    add_header(doc, 'Finance Committee Meeting\nXX-XX-XX', False)

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), 'Times New Roman', size, bold)
    insertHR(title)

    finance = doc.add_paragraph()
    set_font(finance.add_run('Finance Committee Meeting\n'), 'Times New Roman', 12)
    insertHR(finance, 'top')
    set_font(finance.add_run('Date'), 'Times New Roman', 11)
    insertHR(finance)

    parliamentary_officer = doc.add_paragraph()
    set_font(parliamentary_officer.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)
    roles = [('Chair', 'Asst. Tau'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officer, title, role, active_df)
    insertHR(parliamentary_officer)

    set_font(doc.add_paragraph().add_run(f'Call to Order {emDash} Time'), 'Times New Roman', 11, True)
    set_font(doc.add_paragraph().add_run('Asst. Tau Update'), 'Times New Roman', 11, True)
    set_font(doc.add_paragraph().add_run('Business'), 'Times New Roman', 11, True)
    set_font(doc.add_paragraph().add_run(f'Adjournment {emDash} Time'), 'Times New Roman', 11, True)

    doc.add_page_break()

    set_font(doc.add_paragraph().add_run(f'Roster {emDash}'), 'Times New Roman', 11, True)
    set_document_font(doc, font_name='Calibri')

    table_title = doc.add_table(rows=1, cols=1)
    title_cell = table_title.cell(0, 0)
    title_cell.text = 'FINANCE COMMITTEE'
    title_cell.paragraphs[0].style = 'No Spacing'
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    roles = ['Asst. Tau', 'Sigma']
    for role in roles:
        officer_data = active_df[active_df['Current Office'].str.contains(role, na=False)]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            r_index = add_table_row(officers_table, [role, name, 'P'], r_index, center_cols=[2])

    brothers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(roles), na=False)]
    for _, row in others_data.iterrows():
        if r_index >= 6:
            break
        r_index = add_table_row(brothers_table, ['', '', 'P'], r_index, center_cols=[2])

    doc.save(os.path.join(docx_output_dir, 'Finance Committee Outline.docx'))

def create_house_minutes(docx_output_dir, active_df, advisor_df):
    register_element_cls('wp:anchor', CT_Anchor)
    doc = Document()

    paragraph = doc.add_paragraph()
    add_float_picture(paragraph, 'data/AEPKS_BLACK_MALTESE_CROSS.png', width=Inches(3.65), height=Inches(3.95), pos_x=Pt(5), pos_y=Pt(92))

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), size=size, bold=bold)
    insertHR(title)

    meeting = doc.add_paragraph()
    set_font(meeting.add_run('House Meeting\n'), size=12)
    insertHR(meeting, position='top')
    set_font(meeting.add_run('Date'))
    insertHR(meeting, position='bottom')

    parliamentary_officer = doc.add_paragraph()
    set_font(parliamentary_officer.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)
    insertHR(parliamentary_officer, 'top')
    roles = [('Chair', 'Alpha'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officer, title, role, active_df)

    doc.add_page_break()

    agenda = [
        'Call to Order - Time',
        'Chapter Announcements',
        'Alpha Update',
        'Resident Advisor Update',
        'Chapter Advisor Update',
        'Adjournment - Time'
    ]
    for section in agenda:
        set_font(doc.add_paragraph().add_run(section), 'Times New Roman', 11, True)

    doc.add_page_break()

    set_font(doc.add_paragraph().add_run('Roster -'), 'Times New Roman', 11, True)

    officers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    for officer, row in get_officers_from_df(active_df, officers):
        name = f'{row["First Name"]} {row["Last Name"]}'
        r_index = add_table_row(officers_table, [officer, name, 'P', 'P'], r_index, center_cols=[2, 3])

    brothers_table = doc.add_table(rows=1, cols=4)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Brothers'
    hdr_cells[2].text = 'Opening Roll'
    hdr_cells[3].text = 'Closing Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    brothers_data = filter_sorted_brothers(active_df, officers)
    for _, row in brothers_data.iterrows():
        r_index = add_table_row(brothers_table, [row["Last Name"], row["First Name"], 'P', 'P'], r_index, center_cols=[2, 3])

    advisor_table = doc.add_table(rows=1, cols=4)
    set_table_headers(advisor_table, ['Chapter Staff', 'Opening Roll', 'Closing Roll', 'Role'])

    r_index = 0
    for advisor in advisors:
        advisor_data = advisor_df[advisor_df['Current Office'] == advisor]
        for _, row in advisor_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            symbol = 'E' if advisor in ['Resident Advisor', 'Chapter Advisor', 'Asst. Chapter Advisor'] else 'P'
            r_index = add_table_row(advisor_table, [name, symbol, symbol, advisor], r_index, center_cols=[1, 2])

    new_members_table_header = doc.add_table(rows=1, cols=1)
    cell = new_members_table_header.cell(0, 0)
    cell.text = 'NEW MEMBERS'
    cell.paragraphs[0].style = 'No Spacing'
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    new_members_table = doc.add_table(rows=1, cols=4)
    set_table_headers(new_members_table, ['Last Name', 'First Name', 'Opening Roll', 'Closing Roll'])

    r_index = 0
    blank_data = filter_sorted_brothers(active_df, officers)
    for _, _ in blank_data.iterrows():
        r_index = add_table_row(new_members_table, ['', '', 'P', 'P'], r_index, center_cols=[2, 3])

    doc.save(os.path.join(docx_output_dir, 'House Minutes Outline.docx'))

def create_IOC_minutes(docx_output_dir, active_df):
    doc = Document()
    add_header(doc, 'Internal Operation Committee\nXX-XX-XX', False)

    title = doc.add_paragraph()
    for text, size, bold in [('Phi Kappa Sigma\n', 26, True), ('Alpha Epsilon\n', 20, True), ('Meeting Minutes', 14, False)]:
        set_font(title.add_run(text), 'Times New Roman', size, bold)
    insertHR(title)

    ioc = doc.add_paragraph()
    set_font(ioc.add_run('Internal Operations Committee\n'), 'Times New Roman', 12)
    insertHR(ioc, 'top')
    set_font(ioc.add_run('Date'), 'Times New Roman', 11)
    insertHR(ioc)

    parliamentary_officer = doc.add_paragraph()
    set_font(parliamentary_officer.add_run('Parliamentary Officers\n'), 'Times New Roman', 14)
    roles = [('Chair', 'Beta'), ('Secretary', 'Sigma')]
    for title, role in roles:
        add_parliamentary_officers(parliamentary_officer, title, role, active_df)
    insertHR(parliamentary_officer)

    set_font(doc.add_paragraph().add_run(f'Call to Order {emDash} Time'), 'Times New Roman', 11, True)
    add_bullet_section(doc, f'Business {emDash}', ['Theta 1', 'Theta 2', 'Theta 3', 'Business'])
    set_font(doc.add_paragraph().add_run(f'Adjournment {emDash} Time'), 'Times New Roman', 11, True)

    doc.add_page_break()

    set_font(doc.add_paragraph().add_run(f'Roster {emDash}'), 'Times New Roman', 11, True)
    set_document_font(doc, font_name='Calibri')

    table_title = doc.add_table(rows=1, cols=1)
    title_cell = table_title.cell(0, 0)
    title_cell.text = 'INTERNAL OPERATIONS COMMITTEE'
    title_cell.paragraphs[0].style = 'No Spacing'
    title_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    officers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = officers_table.rows[0].cells
    hdr_cells[0].text = 'Officers'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    roles = ['Beta', 'Theta One', 'Theta Two', 'Theta Three', 'Sigma']
    for role in roles:
        officer_data = active_df[active_df['Current Office'].str.contains(role, na=False)]
        for _, row in officer_data.iterrows():
            name = f'{row["First Name"]} {row["Last Name"]}'
            r_index = add_table_row(officers_table, [role, name, 'P'], r_index, center_cols=[2])

    brothers_table = doc.add_table(rows=1, cols=3)
    hdr_cells = brothers_table.rows[0].cells
    hdr_cells[0].text = 'Others'
    hdr_cells[2].text = 'Roll'
    hdr_cells[0].merge(hdr_cells[1])
    for cell in hdr_cells:
        apply_table_header_style(cell)

    r_index = 0
    others_data = active_df[~active_df['Current Office'].str.contains('|'.join(roles), na=False)]
    for _, row in others_data.iterrows():
        if r_index >= 3:
            break
        r_index = add_table_row(brothers_table, ['', '', 'P'], r_index, center_cols=[2])

    doc.save(os.path.join(docx_output_dir, 'IOC Minutes Outline.docx'))
